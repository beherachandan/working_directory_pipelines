"""Convert markdown files to .docx using python-docx, and extract AEO additions."""
import re
from pathlib import Path


def extract_structural_additions(enhanced_md: str, changes_log: str) -> dict:
    """
    Pull structural additions from an enhanced article for appending to a Google Doc.
    Returns dict with keys: faq_section, canonical_answer, citation_replacements.

    citation_replacements: list of {old_text, new_text} for replaceAllText operations.
    """
    result = {"faq_section": "", "canonical_answer": "", "citation_replacements": []}

    # Extract FAQ section from enhanced markdown
    faq_match = re.search(
        r"(##\s*Frequently asked questions.*?)(?=\n##\s|\Z)",
        enhanced_md,
        re.DOTALL | re.IGNORECASE,
    )
    if faq_match:
        result["faq_section"] = faq_match.group(1).strip()

    # Extract citation replacements from changes_log
    # Match lines like: Replaced "X" with "Y" or Replaced 'X' with 'Y'
    if changes_log:
        for line in changes_log.split("\n"):
            m = re.search(
                r'[Rr]eplaced\s+["\u201c]([^"\u201d]+)["\u201d]\s+with\s+["\u201c]([^"\u201d]+)["\u201d]',
                line,
            )
            if m:
                old_text = m.group(1).strip()
                new_text = m.group(2).strip()
                if old_text and new_text and old_text != new_text:
                    result["citation_replacements"].append(
                        {"old_text": old_text, "new_text": new_text, "match_case": False}
                    )

    return result


def build_brand_replacements(brand_issues: list) -> list:
    """
    Convert brand_issues list from article result JSON into replaceAllText operations.
    Returns list of {old_text, new_text, match_case}.
    """
    replacements = []
    for issue in brand_issues:
        rule = issue.get("rule", "")
        instances = issue.get("instances", [])
        fix_text = issue.get("fix", "")

        if rule == "em_dash_in_copy":
            # Replace every em dash with a comma+space
            replacements.append({"old_text": "\u2014", "new_text": ", ", "match_case": True})

        elif rule in ("vocabulary_substitution", "quiz_maker_forbidden_term",
                      "quiz_tool_language", "product_mention"):
            # Extract the exact instance text and apply the fix
            for inst in instances:
                inst_clean = inst.strip("'\"" + "\u201c\u201d")
                if inst_clean and fix_text:
                    # fix_text often says "Replace X with Y" — try to extract Y
                    fix_m = re.search(r"[Rr]eplace\s+.+?\s+with\s+['\u201c]?([^'\u201d,]+)", fix_text)
                    replacement = fix_m.group(1).strip("'\"" + "\u201c\u201d").strip() if fix_m else ""
                    if replacement:
                        replacements.append({
                            "old_text": inst_clean,
                            "new_text": replacement,
                            "match_case": True,
                        })

    # Deduplicate
    seen = set()
    deduped = []
    for r in replacements:
        key = (r["old_text"], r["new_text"])
        if key not in seen:
            seen.add(key)
            deduped.append(r)
    return deduped
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _strip_inline(text: str):
    """Return (text_without_markers, is_bold, is_italic, is_code)."""
    # Strip inline code backticks — just return plain text
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def _add_paragraph_with_inline(doc: Document, text: str, style: str = "Normal"):
    """Add a paragraph handling **bold** and *italic* inline markers."""
    para = doc.add_paragraph(style=style)
    # Split on bold/italic markers
    pattern = re.compile(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)")
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("***") and part.endswith("***"):
            run = para.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            para.add_run(part)
    return para


def markdown_to_docx(md_text: str, output_path: Path, title: str = ""):
    """Convert a markdown string to a .docx file."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Narrow margins for readability
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(90)
        section.right_margin = Pt(90)

    lines = md_text.split("\n")
    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Horizontal rule
        if re.match(r"^---+$", line.strip()):
            doc.add_paragraph("─" * 60, style="Normal").runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue

        # Table detection
        if "|" in line and i + 1 < len(lines) and re.match(r"^\|[-|: ]+\|", lines[i + 1]):
            # Collect all table rows
            table_rows = []
            while i < len(lines) and "|" in lines[i]:
                cols = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                table_rows.append(cols)
                i += 1
            # Skip separator row (index 1)
            if len(table_rows) > 1 and re.match(r"^[-| :]+$", table_rows[1][0].replace("|", "").strip()):
                table_rows.pop(1)
            if table_rows:
                ncols = max(len(r) for r in table_rows)
                tbl = doc.add_table(rows=len(table_rows), cols=ncols)
                tbl.style = "Table Grid"
                for ri, row in enumerate(table_rows):
                    for ci, cell_text in enumerate(row):
                        if ci < ncols:
                            cell = tbl.cell(ri, ci)
                            cell.text = re.sub(r"\*\*(.*?)\*\*", r"\1", cell_text)
                            if ri == 0:
                                for run in cell.paragraphs[0].runs:
                                    run.bold = True
            doc.add_paragraph()
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)

        # Bullet list
        elif line.startswith("* ") or line.startswith("- "):
            _add_paragraph_with_inline(doc, line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            text = re.sub(r"^\d+\. ", "", line)
            _add_paragraph_with_inline(doc, text.strip(), style="List Number")

        # Blank line
        elif line.strip() == "":
            pass  # Skip blank lines (natural paragraph spacing handles it)

        # Normal paragraph
        else:
            _add_paragraph_with_inline(doc, line.strip())

        i += 1

    doc.save(str(output_path))


def convert_all_enhanced(enhanced_dir: Path, docx_dir: Path, article_results: dict | None = None):
    """
    Convert all *-enhanced.md files in enhanced_dir to .docx files in docx_dir.
    Naming: [AEO Updated] {original_name}.docx
    """
    docx_dir.mkdir(parents=True, exist_ok=True)

    md_files = sorted(enhanced_dir.glob("*-enhanced.md"))
    converted = []
    for md_path in md_files:
        # Keep item_num + slug, just swap -enhanced suffix for -aeo-updated
        stem = md_path.stem  # e.g. "52-differentiated-instruction-strategies-enhanced"
        docx_name = stem.replace("-enhanced", "-aeo-updated") + ".docx"
        # e.g. "52-differentiated-instruction-strategies-aeo-updated.docx"
        docx_path = docx_dir / docx_name

        md_text = md_path.read_text(encoding="utf-8")
        try:
            markdown_to_docx(md_text, docx_path, title=docx_name.replace(".docx", ""))
            converted.append((md_path.name, docx_name))
        except Exception as e:
            print(f"  WARN: failed to convert {md_path.name}: {e}")

    return converted
